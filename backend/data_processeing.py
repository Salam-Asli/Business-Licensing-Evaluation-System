
import json
import re
import unicodedata
from pathlib import Path

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import docx
except Exception:
    docx = None

# ---- RTL helpers ----
try:
    from bidi.algorithm import get_display
except Exception:
    get_display = None

# Text utils
HEB_RANGE = r"\u0590-\u05FF"
HEB_WORD = rf"[0-9A-Za-z_{HEB_RANGE}]"
heb_run = re.compile(fr"[{HEB_RANGE}]+")

## function to extarct data from pdf
def extract_text_from_pdf(path):
    if not pdfplumber:
        raise RuntimeError("pdfplumber is not installed")
    text = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text.append(page.extract_text() or "")
    return "\n".join(text)


### function to extract data from docx
def extract_text_from_docx(path):
    if not docx:
        raise RuntimeError("docx is not installed")
        text=[]
    d = docx.Document(path)
    return "\n".join([p.text for p in d.paragraphs if p.text.strip()])

# for clean strings only chars
def normalize_text(s: str) -> str:
    if s is None:
        return ""
    s = unicodedata.normalize("NFKC", s)
    s = s.replace("\u200f", "").replace("\u200e", "")
    return "\n".join(" ".join(line.split()) for line in s.splitlines()).strip()


def fallback_fix_rtl(s: str) -> str:
    s = normalize_text(s)
    return heb_run.sub(lambda m: m.group(0)[::-1], s)

def fix_rtl(s: str) -> str:
    s = normalize_text(s)
    if get_display:
        try:
            return get_display(s)
        except Exception:
            pass
    return fallback_fix_rtl(s)

def hb(token: str) -> str:
    return rf"(?<!{HEB_WORD}){token}(?!{HEB_WORD})"

HEADING_PATTERN = re.compile(
    r"^(?:\d+(?:\.\d+)+\s+.+|[\u05D0-\u05EA0-9\s\(\)\-]{3,}:?)$"
)

## split text by headings
def extract_text_from_docx(path: Path) -> str:
    if not docx:
        raise RuntimeError("docx is not installed")
    d = docx.Document(path)
    # keep only non-empty lines
    return "\n".join([p.text for p in d.paragraphs if p.text.strip()])

# -------------------- Sectioning --------------------
HEADING_PATTERN = re.compile(
    r"^(?:\d+(?:\.\d+)+\s+.+|[\u05D0-\u05EA0-9\s\(\)\-]{3,}:?)$"
)

def split_by_headings(full_text: str):
    """
    Split the raw text by headings. Keep BOTH:
      - 'raw'  : original text for matching (avoids over-fixing)
      - 'text' : display text with RTL repaired (for readable output)
    """
    lines = full_text.splitlines()
    sections = []

    current_heading = "Introduction"
    current_text = []

    for line in lines:
        raw_line = line
        line = line.strip()
        if not line:
            continue

        if HEADING_PATTERN.match(line) and len(line) <= 80:
            if current_text:
                raw_block = "\n".join(current_text)
                sections.append({
                    "heading": fix_rtl(current_heading),
                    "text": fix_rtl(raw_block),  # pretty version
                    "raw": raw_block             # original for matching
                })
                current_text = []
            current_heading = line  # keep raw; fix on output only
        else:
            current_text.append(raw_line)

    if current_text:
        raw_block = "\n".join(current_text)
        sections.append({
            "heading": fix_rtl(current_heading),
            "text": fix_rtl(raw_block),
            "raw": raw_block
        })
    return sections

def map_business_requierments(sections):
    ## map each section by keyword
    ## patterns i used:
    patterns = {
        "gas_usage": rf"({hb('גז')}|מערכת\s+גז|גפ\"מ|תשתית\s+גז)",
        "meat_service": rf"({hb('בשר')}|{hb('עופות')}|{hb('דגים')}|אחסון\s+בשר|הפרדה\s+(?:בשרי|בשר\s*.*?חלבי|דגים))",
        "seating": r"(מקומות\s+ישיבה|קיבולת(?:\s+קהל)?\s+מקסימלית|תפוסה\s+מותרת)"
    }
    compiled = {k: re.compile(v, re.I | re.M) for k, v in patterns.items()}
    mapped = []   ## empty list
    for sec in sections:
        # match against both variants
        candidates = [sec.get("raw", ""), sec.get("text", "")]
        scores = {k: 0 for k in compiled}
        for cand in candidates:
            for key, rx in compiled.items():
                scores[key] += len(rx.findall(cand))

        best_key, best_score = max(scores.items(), key=lambda kv: kv[1])
        others = sum(v for k, v in scores.items() if k != best_key)

        # accept only strong, unambiguous matches
        if best_score > 0 and best_score > others:
            mapped.append({
                "key": best_key,
                "heading": sec["heading"],
                "context": sec["text"]  # readable (fixed) context
            })
    return mapped


def main():
    pdf_file = Path("18-07-2022_4.2A.pdf")
    docx_file = Path("18-07-2022_4.2A.docx")

    if pdf_file.exists() and pdf_file.suffix.lower() == ".pdf":
        raw_text = extract_text_from_pdf(pdf_file)
        source_file = pdf_file
    elif docx_file.exists() and docx_file.suffix.lower() == ".docx":
        raw_text = extract_text_from_docx(docx_file)
        source_file = docx_file
    else:
        raise FileNotFoundError("No input file found in data/")


    sections = split_by_headings(raw_text)
    parsed = map_business_requierments(sections)
    output_path = Path("processed.json")
    output_path.write_text(
        json.dumps(
            {
                "source_file": str(source_file),  # ✅ make sure it's string
                "parsed": parsed
            },
            ensure_ascii=False,
            indent=2
        ),
        encoding="utf-8"
    )


if __name__ == "__main__":
    main()